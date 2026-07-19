import csv
import json
from io import BytesIO, StringIO
from typing import Any, Dict
from docx import Document


# ----------------------------
# CSV export
# ----------------------------

 
def to_csv_bytes(report: Dict[str, Any]) -> bytes:
    issue_inventory = {item["id"]: item for item in report.get("issue_inventory", [])}
    issue_results = {item["issue_id"]: item for item in report.get("issue_results", [])}
    rows = []
    for issue_id, issue in issue_inventory.items():
        result = issue_results.get(issue_id, {})
        rows.append(
            {
                "issue_id": issue_id,
                "title": issue.get("title", ""),
                "target_rule": issue.get("target_rule", ""),
                "inventory_summary": issue.get("summary", ""),
                "verdict": result.get("verdict", ""),
                "rationale": result.get("rationale", ""),
                "evidence": " | ".join(result.get("evidence", []) or []),
                "suggested_fix": result.get("suggested_fix", ""),
            }
        )

 
    buffer = StringIO()
    writer = csv.DictWriter(buffer, fieldnames=list(rows[0].keys()) if rows else [
        "issue_id", "title", "target_rule", "inventory_summary",
        "verdict", "rationale", "evidence", "suggested_fix"
    ])
    writer.writeheader()
    writer.writerows(rows)
    return buffer.getvalue().encode("utf-8")

 
# ----------------------------
# Word export
# ----------------------------

def to_word_bytes(report: Dict[str, Any]) -> bytes:
    doc = Document()
    doc.add_heading("Compliance Verification Report", level=1)
    doc.add_paragraph(f"Final verdict: {report.get('verdict', '')}")
    if report.get("summary"):
        doc.add_paragraph(f"Summary: {report.get('summary', '')}")
    if report.get("explanation"):
        doc.add_paragraph(f"Explanation: {report.get('explanation', '')}")

    doc.add_heading("Issue Results", level=2)
    for item in report.get("issue_results", []):
        doc.add_heading(f"{item.get('issue_id', '')} — {item.get('verdict', '')}", level=3)
        doc.add_paragraph(f"Rationale: {item.get('rationale', '')}")
        if item.get("evidence"):
            doc.add_paragraph("Evidence: " + " | ".join(item.get("evidence", [])))
        if item.get("suggested_fix"):
            doc.add_paragraph(f"Suggested fix: {item.get('suggested_fix', '')}")

    doc.add_heading("Evidence Map", level=2)
    for issue_id, evidence in report.get("evidence_map", {}).items():
        doc.add_paragraph(f"{issue_id}: " + " | ".join(evidence or []))


    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()